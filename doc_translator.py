import os
from docx import Document
from gpt_translator import GPTTranslator
from concurrent.futures import ThreadPoolExecutor, wait
import json
import shutil
import time


class DocTranslator:
    def __init__(self, num_threads=16):
        self.num_threads = num_threads
        terminology = self.__load_terminology("terminology.txt")
        print(f"术语表: {json.dumps(terminology, indent=4)}")

        self.translator = GPTTranslator(terminology)

    def __load_terminology(self, file_path):
        data = {}
        with open(file_path, "r", encoding="utf-8") as file:
            for line in file:
                # 去除每行的空白符
                line = line.strip()
                if line:
                    # 分割成键值对
                    key, value = line.split("=", 1)
                    # 去除键和值两端的空白符
                    key = key.strip()
                    value = value.strip()
                    # 添加到字典中
                    data[key] = value
        return data

    def __translate_text(self, text):
        return self.translator.translate(text)
        # return text

    def __copy_file(self, source_file, dest_file):
        try:
            shutil.copyfile(source_file, dest_file)
        except Exception as e:
            print(f"文件复制失败: {e}")

    def __translate_paragraph(self, index, paragraph):
        print(f"############# 正在处理 {self.part} 第 {index} 个段落 #############")
        if paragraph.text.strip():
            translated_text = self.__translate_text(paragraph.text)
            print(f"原文: {paragraph.text}")
            print(f"译文: {translated_text}")
            paragraph.text = paragraph.text.replace(
                paragraph.text.strip(), translated_text.strip()
            )

    def __translate_paragraphs(self, paragraphs):
        paragraphs = [paragraph for paragraph in paragraphs if paragraph.text.strip()]
        # 只处理前30个
        # paragraphs = paragraphs[:30]

        for index, item in enumerate(paragraphs):
            self.__translate_paragraph(index, item)

    def __translate_tables(self, tables):
        paragraphs = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        self.__translate_paragraphs(paragraphs)

    def __translate_headers(self, headers):
        paragraphs = []
        for header in headers:
            paragraphs.extend(header.paragraphs)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
        self.__translate_paragraphs(paragraphs)

    def __translate_footers(self, footers):
        paragraphs = []
        for footer in footers:
            paragraphs.extend(footer.paragraphs)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.extend(cell.paragraphs)
        self.__translate_paragraphs(paragraphs)

    def translate(self, file):
        start = time.time()

        # 复制文件
        translated_file = file.replace(".docx", "-translated.docx")
        self.__copy_file(file, translated_file)

        doc = Document(translated_file)

        # 翻译页眉
        # headers = [section.header for section in doc.sections]
        self.part = "页眉"
        self.__translate_headers([doc.sections[0].header])

        # 翻译页脚
        # footers = [section.footer for section in doc.sections]
        self.part = "页脚"
        self.__translate_footers([doc.sections[0].footer])

        # 翻译shapes
        # self.__translate_paragraphs(doc.inline_shapes)
        self.part = "shapes"
        for shape in doc.inline_shapes:
            print(shape)

        # 翻译段落
        self.part = "段落"
        self.__translate_paragraphs(doc.paragraphs)

        # 翻译表格
        self.part = "表格"
        self.__translate_tables(doc.tables)

        # 保存文件
        doc.save(translated_file)

        end = time.time()
        print(f"翻译完成，用时: {end - start:.2f} 秒")


if __name__ == "__main__":
    translator = DocTranslator()
    translator.translate("NSR_T103508-7_PH-42754_DRF dog.docx")
