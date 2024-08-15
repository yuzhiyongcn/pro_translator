import os
from docx import Document
from gpt_translator import GPTTranslator
import json
import shutil
import time
import terms_persister as persister
import string_utils as su
import file_utils as fu

MAX_TOKEN_LIMIT = 2500


class DocTranslator:
    def __init__(self, to_CN=False):
        self.to_CN = to_CN

        # terms_file = r"input\terms_en2zh.xlsx" if to_CN else r"input\terms_zh2en.xlsx"
        # self.terms = persister.read_from_excel(terms_file)
        # 不适用术语表
        self.terms = {}

        additional_terms = (
            r"input\additional_terms_en2zh.json"
            if to_CN
            else r"input\additional_terms_zh2en.json"
        )
        self.terms.update(fu.read_json_file(additional_terms))

        print(f"术语表单词数量: {len(self.terms)}")

        self.translator = GPTTranslator(to_CN)
        self.excludes = ["F", "Y", "N", "-", "+"]
        self.debug_mode = True
        self.last_to_translate_amount = 0

    def _translate_text(self, text):
        if text in self.excludes:
            return text
        if text.replace(" ", "").isdigit():
            return text

        return self.translator.translate(text)

    def _copy_file(self, source_file, dest_file):
        try:
            shutil.copyfile(source_file, dest_file)
        except Exception as e:
            print(f"文件复制失败: {e}")

    def __load_tables(self, tables):
        paragraphs = []
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        return paragraphs

    def __load_headers_or_footers(self, headers_or_footers):
        paragraphs = []
        for header_or_footer in headers_or_footers:
            paragraphs.extend(header_or_footer.paragraphs)
            paragraphs.extend(self.__load_tables(header_or_footer.tables))
        return paragraphs

    def __is_text_skipped(self, text):
        text = text.replace(" ", "")
        return text in self.excludes or text.isdigit()

    def __preprocess(self, text_list):
        preprocessed_text_list = []
        for text in text_list:
            for term_key, term_value in self.terms.items():
                if term_key in text:
                    text = text.replace(term_key, term_value)

            preprocessed_text_list.append(text)

        return preprocessed_text_list

    def __is_translate_required(self, text):
        if not self.__is_text_skipped(text):
            if self.to_CN and not su.has_chinese(text):  # 翻译成中文，只翻译英文文本
                return True
            elif not self.to_CN and su.has_chinese(text):  # 翻译成英文，只翻译中文文本
                return True
        return False

    # 递归方法, 将待翻译文本列表翻译成中文, 并将翻译结果存入translated_dict
    # to_translate_list数量为0时, 翻译结束
    def __translate_list(self, to_translate_list, translated_dict):
        if len(to_translate_list) == 0:
            return

        # 分割待翻译文本
        to_translate = []
        while (
            len("[#]".join(to_translate)) < MAX_TOKEN_LIMIT
            and len(to_translate_list) > 0
        ):
            to_translate.append(to_translate_list.pop(0))

        # 翻译
        preprocessed = self.__preprocess(to_translate)
        to_translate_text = "[#]".join(preprocessed)
        translated_text = self.translator.translate(to_translate_text)
        # print(f"---------------原文----------")
        # print(to_translate_text)
        # print(f"---------------翻译结果----------")
        # print(translated_text)
        # print(f"---------------翻译结果结束----------")
        translated_list = translated_text.split("[#]")
        for i, text in enumerate(to_translate):
            translated_dict[text] = translated_list[i]

        print(f"本次翻译数量: {len(to_translate)}")
        print(
            f"已翻译: {len(translated_dict)} / {len(translated_dict)+len(to_translate_list)}"
        )

        # 递归
        self.__translate_list(to_translate_list, translated_dict)

        pass

    def translate(self, file):
        start = time.time()
        self.debug_file_path = file.replace(".docx", f"-debug-{time.time()}.txt")
        self.debug_file_path = "output\\" + self.debug_file_path

        # 复制文件
        translated_file = file.replace(".docx", "-translated.docx")
        self._copy_file(file, translated_file)

        doc = Document(translated_file)

        paragraphs = []
        # 读取所有段落
        headers = [section.header for section in doc.sections]
        paragraphs.extend(self.__load_headers_or_footers(headers))
        footers = [section.footer for section in doc.sections]
        paragraphs.extend(self.__load_headers_or_footers(footers))
        # shapes = []
        # for shape in doc.inline_shapes:
        #     shapes.append(shape)
        # paragraphs.extend(shapes)
        paragraphs.extend(doc.paragraphs)
        paragraphs.extend(self.__load_tables(doc.tables))

        # 获取所有段落文本
        texts = set()
        for paragraph in paragraphs:
            texts.add(paragraph.text)

        # 分离中英文文本
        to_translate_texts = []
        for text in texts:
            text = text.strip()
            if self.__is_translate_required(text):
                to_translate_texts.append(text)

        print(f"待翻译文本数量: {len(to_translate_texts)}")

        # 只翻译前30个
        # to_translate_texts = to_translate_texts[:30]

        if len(to_translate_texts) == 0:
            print("------------翻译已经全部完成------------")
            # 删除translated_file
            os.remove(translated_file)
            return

        # 翻译
        # 初始化翻译结果
        translated_dict = {}
        self.__translate_list(to_translate_texts, translated_dict)

        # 写入log
        if self.debug_mode:
            with open(self.debug_file_path, "w", encoding="utf-8") as debug_file:
                for key, value in translated_dict.items():
                    debug_file.write(f"原文: {key}\n\n")
                    debug_file.write(f"译文: {value}\n")
                    debug_file.write(
                        "------------------------------------------------------------------------------------\n"
                    )

        # 更新段落文本
        for paragraph in paragraphs:
            key = paragraph.text.strip()
            if key in translated_dict:
                original_text = paragraph.text
                translated_text = translated_dict[key]
                # 双语翻译
                paragraph.text = paragraph.text.replace(
                    key, original_text + "\n" + translated_text
                )
                # 单语翻译
                # paragraph.text = paragraph.text.replace(key, translated_text)
        pass

        # 保存文件
        doc.save(translated_file)

        end = time.time()
        print(f"翻译完成，用时: {end - start:.2f} 秒")


if __name__ == "__main__":
    translator = DocTranslator(False)
    translator.translate(
        "苏州华测_非临床安全性评价研究方案及报价_化药1类- 双语-0814.docx"
    )
